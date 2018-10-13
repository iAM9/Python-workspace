#!/usr/bin/env python

#################################
import sys
import os
import shutil

import Tkinter
from tkMessageBox import *
from tkFileDialog import *

from docx import Document
#################################

def compile(absPath):
    # Arrays to hold list of files and their names
    listOfFilePaths = []
    listOfFileNames = []


    for path, subdirs, files in os.walk(absPath):
        for name in files:
            if name.endswith(".txt") or name.endswith(".doc") or name.endswith(".docx"):
                filePath = os.path.join(path, name)
                listOfFileNames.append(name)
                listOfFilePaths.append(filePath)

    print ("\n".join(listOfFilePaths))
    print "\n"
    print ("\n".join(listOfFileNames))

    # Create new word document
    document = Document()
    document.add_heading('Ideas', 0)

    # Open text files one by one, read them and write them to a word ".docx" file
    for i in xrange(0, len(listOfFileNames)):

        # If it's a text file
        if listOfFileNames[i].endswith(".txt"):
            opentxt = open(listOfFilePaths[i], 'r')
            lines = opentxt.readlines()
            opentxt.close()
            document.add_heading(listOfFileNames[i], level = 3)
            for ls in lines:
                document.add_paragraph(ls)
            document.add_paragraph()

        # If it's a word 'doc' or 'docx' file    
        elif listOfFileNames[i].endswith(".docx"):
            doc = Document(listOfFilePaths[i])
            document.add_heading(listOfFileNames[i], level = 3)
            for p in doc.paragraphs:
                document.add_paragraph(p.text)
            document.add_paragraph()

    finalName = absPath + "//" + "Compiled ideas.docx"
    print (finalName)
    document.save(finalName)
    print ("Document saved successfully!")
    showinfo("Success", "Compiling finished")

##############################################
def cmd_about():
    showinfo("About", "\tCompilinator\n\nProgrammed by Ibrahim Asif Mirza\n\nIdea by: Nadir H. Shah")

def cmd_selectFolder():
    absPath =  askdirectory()
    if absPath == '':
        print "Nothing selected"
    else:
        print absPath
        compile(absPath)

top = Tkinter.Tk()
top.title("Compilinator")
Tkinter.Label(top, text="== Compile your ideas ==").pack()
Tkinter.Button(top, text='Select Folder', command=cmd_selectFolder).pack()
Tkinter.Button(top, text='About', command=cmd_about).pack()

top.mainloop()


##############################################
